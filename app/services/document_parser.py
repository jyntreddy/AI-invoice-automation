import os
from typing import Optional, Dict, Any, List
import PyPDF2
import docx
import pandas as pd
from pathlib import Path

from app.core.logging import get_logger
from app.core.config import settings

logger = get_logger()


class DocumentParser:
    """Service for parsing various document formats."""
    
    @staticmethod
    def parse_document(file_path: str) -> Dict[str, Any]:
        """Parse document based on file extension."""
        file_ext = Path(file_path).suffix.lower()
        
        parsers = {
            '.pdf': DocumentParser.parse_pdf,
            '.docx': DocumentParser.parse_docx,
            '.doc': DocumentParser.parse_docx,
            '.csv': DocumentParser.parse_csv,
            '.xlsx': DocumentParser.parse_excel,
            '.xls': DocumentParser.parse_excel,
        }
        
        parser = parsers.get(file_ext)
        if not parser:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        try:
            result = parser(file_path)
            logger.info(f"Successfully parsed document: {file_path}")
            return result
        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """Parse PDF document."""
        text_content = []
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title'),
                        'author': pdf_reader.metadata.get('/Author'),
                        'subject': pdf_reader.metadata.get('/Subject'),
                        'creator': pdf_reader.metadata.get('/Creator'),
                    }
                
                # Extract text from all pages
                num_pages = len(pdf_reader.pages)
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                
            full_text = "\n".join(text_content)
            
            return {
                'text': full_text,
                'num_pages': num_pages,
                'metadata': metadata,
                'file_type': 'pdf'
            }
            
        except Exception as e:
            raise Exception(f"PDF parsing error: {str(e)}")
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """Parse DOCX document."""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            full_text = "\n".join(paragraphs)
            
            # Extract core properties
            core_properties = doc.core_properties
            metadata = {
                'title': core_properties.title,
                'author': core_properties.author,
                'subject': core_properties.subject,
                'created': core_properties.created,
                'modified': core_properties.modified,
            }
            
            return {
                'text': full_text,
                'paragraphs': paragraphs,
                'tables': tables_data,
                'metadata': metadata,
                'file_type': 'docx'
            }
            
        except Exception as e:
            raise Exception(f"DOCX parsing error: {str(e)}")
    
    @staticmethod
    def parse_csv(file_path: str) -> Dict[str, Any]:
        """Parse CSV file."""
        try:
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            text = df.to_string(index=False)
            
            return {
                'text': text,
                'dataframe': df.to_dict(orient='records'),
                'columns': df.columns.tolist(),
                'num_rows': len(df),
                'metadata': {
                    'shape': df.shape,
                    'columns': df.columns.tolist(),
                },
                'file_type': 'csv'
            }
            
        except Exception as e:
            raise Exception(f"CSV parsing error: {str(e)}")
    
    @staticmethod
    def parse_excel(file_path: str) -> Dict[str, Any]:
        """Parse Excel file."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            all_text = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data[sheet_name] = {
                    'data': df.to_dict(orient='records'),
                    'columns': df.columns.tolist(),
                    'num_rows': len(df)
                }
                all_text.append(f"\n=== Sheet: {sheet_name} ===\n")
                all_text.append(df.to_string(index=False))
            
            full_text = "\n".join(all_text)
            
            return {
                'text': full_text,
                'sheets': sheets_data,
                'sheet_names': excel_file.sheet_names,
                'metadata': {
                    'num_sheets': len(excel_file.sheet_names),
                },
                'file_type': 'xlsx'
            }
            
        except Exception as e:
            raise Exception(f"Excel parsing error: {str(e)}")
    
    @staticmethod
    def extract_invoice_fields(text: str) -> Dict[str, Any]:
        """Extract common invoice fields from text using pattern matching."""
        import re
        from datetime import datetime
        
        extracted = {}
        
        # Invoice number patterns
        invoice_patterns = [
            r'invoice\s*#?:?\s*([A-Z0-9-]+)',
            r'inv\s*#?:?\s*([A-Z0-9-]+)',
            r'invoice\s*number:?\s*([A-Z0-9-]+)',
        ]
        for pattern in invoice_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                extracted['invoice_number'] = match.group(1)
                break
        
        # Date patterns
        date_patterns = [
            r'date:?\s*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            r'invoice\s*date:?\s*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
        ]
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                extracted['invoice_date'] = match.group(1)
                break
        
        # Amount patterns
        amount_patterns = [
            r'total:?\s*\$?\s*([\d,]+\.?\d*)',
            r'amount:?\s*\$?\s*([\d,]+\.?\d*)',
            r'grand\s*total:?\s*\$?\s*([\d,]+\.?\d*)',
        ]
        for pattern in amount_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                amount_str = match.group(1).replace(',', '')
                try:
                    extracted['total'] = float(amount_str)
                except ValueError:
                    pass
                break
        
        # Vendor/Company name (usually at the top)
        lines = text.split('\n')
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 3 and not re.match(r'^[\d\s\-/]+$', line):
                extracted['vendor_name'] = line
                break
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            extracted['vendor_email'] = emails[0]
        
        return extracted


# Singleton instance
document_parser = DocumentParser()
